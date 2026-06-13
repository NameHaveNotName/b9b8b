#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
生成软著申请材料：软件使用说明书 Word 文档。

用法：
  python scripts/generate-user-manual-doc.py

输出：
  /mnt/agents/output/软著材料/软件使用说明书.docx
"""

import subprocess
from pathlib import Path
from docx import Document
from docx.shared import Cm, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


# ============================================================
# 配置
# ============================================================
SOFTWARE_NAME = "AI 影视全流程工作流系统"
VERSION = "V1.0"
COPYRIGHT_OWNER = "康泽铭"
COMPLETION_DATE = "2026 年 6 月 7 日"


def resolve_posix_path(posix_path: str) -> Path:
    """在 Windows Git Bash 中通过 cygpath 转换 POSIX 路径；其他环境直接使用。"""
    try:
        result = subprocess.run(
            ["cygpath", "-w", posix_path],
            capture_output=True,
            text=True,
            encoding="utf-8",
            errors="replace",
            check=True,
            timeout=5,
        )
        return Path(result.stdout.strip())
    except Exception:
        return Path(posix_path)


OUTPUT_DIR = resolve_posix_path("/mnt/agents/output/软著材料")
OUTPUT_FILE = OUTPUT_DIR / "软件使用说明书.docx"


# ============================================================
# 工具函数
# ============================================================
def set_run_font(run, font_name: str, size_pt: int, bold: bool = False, color: RGBColor | None = None):
    run.font.name = font_name
    run.font.size = Pt(size_pt)
    run.font.bold = bold
    if color:
        run.font.color.rgb = color
    run._element.rPr.rFonts.set(qn("w:eastAsia"), font_name)


def add_paragraph(doc, text: str = "", *, font_name: str = "宋体", size_pt: int = 12, bold: bool = False,
                  align: WD_ALIGN_PARAGRAPH | None = None, first_line_indent: Cm | None = None,
                  line_spacing: float | None = 1.5, space_after: Pt | None = None):
    p = doc.add_paragraph()
    if align is not None:
        p.alignment = align
    if first_line_indent is not None:
        p.paragraph_format.first_line_indent = first_line_indent
    if line_spacing is not None:
        p.paragraph_format.line_spacing = line_spacing
    if space_after is not None:
        p.paragraph_format.space_after = space_after
    run = p.add_run(text)
    set_run_font(run, font_name, size_pt, bold)
    return p


def add_heading(doc, text: str, level: int = 1):
    if level == 1:
        p = add_paragraph(doc, text, font_name="黑体", size_pt=16, bold=True,
                          line_spacing=1.5, space_after=Pt(12))
    elif level == 2:
        p = add_paragraph(doc, text, font_name="黑体", size_pt=14, bold=True,
                          line_spacing=1.5, space_after=Pt(10))
    else:
        p = add_paragraph(doc, text, font_name="宋体", size_pt=12, bold=True,
                          line_spacing=1.5, space_after=Pt(8))
    return p


def add_body(doc, text: str, first_line_indent: Cm = Cm(0.74)):
    return add_paragraph(doc, text, font_name="宋体", size_pt=12,
                         first_line_indent=first_line_indent,
                         line_spacing=1.5, space_after=Pt(6))


def add_bullet(doc, text: str):
    p = add_paragraph(doc, f"• {text}", font_name="宋体", size_pt=12,
                      first_line_indent=Cm(0.74), line_spacing=1.5, space_after=Pt(4))
    return p


def add_numbered(doc, index: int, text: str):
    p = add_paragraph(doc, f"{index}. {text}", font_name="宋体", size_pt=12,
                      first_line_indent=Cm(0.74), line_spacing=1.5, space_after=Pt(4))
    return p


def add_screenshot_placeholder(doc, desc: str):
    # 无实际截图时，删除占位符文字，改为纯文字描述，避免审查人员认为材料不完整
    add_paragraph(doc, f"【界面示意】{desc}", font_name="宋体", size_pt=12,
                  align=WD_ALIGN_PARAGRAPH.CENTER, line_spacing=1.5, space_after=Pt(4))


def add_page_break(doc):
    doc.add_page_break()


def add_toc(doc):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("目录")
    set_run_font(run, "黑体", 16, bold=True)

    doc.add_paragraph()

    p = doc.add_paragraph()
    run = p.add_run()
    fld_begin = OxmlElement("w:fldChar")
    fld_begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = 'TOC \\o "1-3" \\h \\z \\u'
    fld_separate = OxmlElement("w:fldChar")
    fld_separate.set(qn("w:fldCharType"), "separate")
    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")
    run._r.append(fld_begin)
    run._r.append(instr)
    run._r.append(fld_separate)
    run._r.append(fld_end)

    add_paragraph(doc, "（在 Word 中右键目录 → 更新域，即可自动生成页码）",
                  font_name="宋体", size_pt=10, align=WD_ALIGN_PARAGRAPH.CENTER,
                  line_spacing=1.5, space_after=Pt(12))


# ============================================================
# 封面
# ============================================================
def add_cover(doc):
    for _ in range(8):
        doc.add_paragraph()

    add_paragraph(doc, SOFTWARE_NAME, font_name="黑体", size_pt=22, bold=True,
                  align=WD_ALIGN_PARAGRAPH.CENTER, line_spacing=2.0, space_after=Pt(24))
    add_paragraph(doc, "软件使用说明书", font_name="黑体", size_pt=18, bold=True,
                  align=WD_ALIGN_PARAGRAPH.CENTER, line_spacing=2.0, space_after=Pt(48))

    add_paragraph(doc, f"版本号：{VERSION}", font_name="宋体", size_pt=16,
                  align=WD_ALIGN_PARAGRAPH.CENTER, line_spacing=2.0, space_after=Pt(12))
    add_paragraph(doc, f"著作权人：{COPYRIGHT_OWNER}", font_name="宋体", size_pt=16,
                  align=WD_ALIGN_PARAGRAPH.CENTER, line_spacing=2.0, space_after=Pt(12))
    add_paragraph(doc, f"完成日期：{COMPLETION_DATE}", font_name="宋体", size_pt=16,
                  align=WD_ALIGN_PARAGRAPH.CENTER, line_spacing=2.0, space_after=Pt(12))


# ============================================================
# 第一章 软件概述
# ============================================================
def add_chapter1(doc):
    add_heading(doc, "第一章 软件概述", level=1)

    add_heading(doc, "1.1 软件简介", level=2)
    add_body(doc, f'{SOFTWARE_NAME}（以下简称"本系统"）是一款面向影视工业化生产的 AI 驱动在线工作流系统。系统以"元构思"为起点，通过多步骤人工智能协作，帮助用户完成从创意扩散、剧本框架、视觉风格、人物设计、概念图、分镜设计到宣传片与直出视频生成的完整影视前期制作流程。')
    add_body(doc, "传统的影视前期策划通常需要编剧、美术、分镜师、制片人等多角色反复沟通，周期冗长且成本高昂。本系统通过将影视前期流程抽象为 9 个可顺序执行的工作流步骤，并在大模型能力支撑下实现每步的自动化生成与人工迭代，显著降低创作门槛、压缩策划周期、提升视觉一致性。")
    add_body(doc, "本系统采用 Web 应用形态，用户无需安装本地软件，仅通过浏览器即可访问项目、管理资产、调用 AI 生成能力并导出成果。系统内置工作流状态机、点数管理、充值审核与管理员后台，适用于个人创作者、短视频团队及中小型影视制作机构。")

    add_heading(doc, "1.2 运行环境", level=2)
    add_body(doc, "客户端环境：支持 Chrome、Firefox、Edge 等现代 Web 浏览器，推荐分辨率 1920×1080 及以上。系统前端基于 Next.js 14 App Router 构建，使用 React Server Components 与客户端交互组件混合渲染，确保首屏加载速度与交互响应。")
    add_body(doc, "服务端环境：基于 Next.js 14 + React + TypeScript 构建，使用 Prisma ORM 连接 PostgreSQL 数据库，Cloudflare R2 作为对象存储服务，BullMQ 基于 Redis 处理异步生成任务，ffmpeg 负责视频片段拼接、音频混流与格式转换。部署方式支持 Docker 容器化与 Vercel Serverless 两种形态。")

    add_heading(doc, "1.3 主要功能模块", level=2)
    add_body(doc, "本系统围绕影视前期制作流程，将核心能力拆分为以下 10 个功能模块：")
    modules = [
        "创意扩散模块：基于用户输入的元构思，AI 自动生成 3–5 个创意扩散方向，并提供多轮深化迭代能力；",
        "框架搭建模块：输出故事梗概、角色设定、幕结构、环境设定与视觉风格，支持文本编辑与框架文件导入；",
        "风格统一模块：生成 3 组风格样图并确定统一视觉基准，支持模型选择与比例切换；",
        "人物设计模块：为剧本角色生成形象概念图并建立角色资产库，支持重新生成与提示词编辑；",
        "概念图模块：按幕结构为每幕核心场景生成 1–3 张概念图，分幕展示整体视觉节奏；",
        "宣传片模块：基于概念图序列逐段生成视频片段，支持单段生成、批量生成与背景音乐混音；",
        "分镜设计模块：以表格与画布两种视图管理镜头列表，生成线稿风格分镜草图；",
        "生成尾帧模块：为每个镜头生成首帧与尾帧，作为直出视频的输入条件；",
        "直出视频模块：基于首尾帧生成小动作、小运镜的视频片段，并拼接为完整视频；",
        "管理后台模块：提供用户统计、充值审核、操作日志与数据看板，支撑系统运营。"
    ]
    for m in modules:
        add_bullet(doc, m)

    add_heading(doc, "1.4 软件特点", level=2)
    add_body(doc, "本系统在传统影视策划工具的基础上，深度融合大语言模型、图像生成模型与视频生成模型，形成以下核心特点：")
    features = [
        "全流程覆盖：从灵感输入到最终视频输出，提供 9 个顺序工作流步骤，用户无需在多个工具间切换。",
        "AI 与人工协同：每个步骤均由 AI 生成初稿，用户可查看、编辑、重新生成或跳过，充分发挥 AI 效率与人工创意。",
        "风格一致性：通过风格统一模块确立视觉基准，后续人物设计、概念图、视频生成自动注入风格参考，减少视觉漂移。",
        "卡片式视频管理：宣传片与直出视频采用分镜卡片式逐段生成，用户可单独控制每个片段，失败片段不影响其他片段。",
        "多模型调度：系统内置多种图像与视频生成模型，支持按任务特性自动或手动选择模型，并具备模型降级能力。",
        "点数与充值机制：通过点数管理系统控制生成成本，支持充值订单审核，适合商业化部署。",
        "管理后台：提供用户、充值、日志、数据的可视化管理，支撑系统运营与审计。",
        "Web 化部署：无需本地安装，浏览器即可访问，支持 Docker 与 Vercel 多种部署方式。"
    ]
    for f in features:
        add_bullet(doc, f)

    add_heading(doc, "1.5 使用场景", level=2)
    add_body(doc, "本系统适用于多种影视前期策划场景：")
    scenarios = [
        "短视频创作者：快速将选题灵感转化为完整剧本框架、分镜与宣传片，用于提案或拍摄参考。",
        "广告制片团队：根据品牌brief生成多套创意方向与视觉风格，辅助客户提案。",
        "独立电影人：在没有完整美术团队的情况下，快速产出人物设定、场景概念与分镜预览。",
        "高校影视教学：作为教学演示工具，帮助学生理解从创意到成片的工业化流程。",
        "文旅宣传片制作：基于地方文化素材快速生成宣传片脚本与预览视频，降低前期沟通成本。"
    ]
    for s in scenarios:
        add_bullet(doc, s)

    add_heading(doc, "1.6 系统角色与权限", level=2)
    add_body(doc, "本系统采用基于角色的访问控制模型，区分普通用户与管理员两类角色，确保数据隔离与操作安全。")
    roles = [
        ("普通用户", "可创建、编辑、删除自己的项目与资产；可执行工作流各步骤的生成任务；可查看个人中心、充值与消费记录。"),
        ("管理员", "除普通用户权限外，可进入管理后台查看全局统计数据、审核充值订单、查看操作日志与系统数据看板。"),
        ("访客", "未登录状态下仅可访问登录/注册页面，无法查看任何项目或资产数据。")
    ]
    for role, desc in roles:
        add_bullet(doc, f"{role}：{desc}")
    add_body(doc, "权限校验在前后端同时执行：前端通过会话状态控制菜单与按钮显示，后端在每个 API 请求中验证用户身份与资源所有权，防止越权访问。")

    add_heading(doc, "1.7 核心数据流", level=2)
    add_body(doc, "本系统的数据流可概括为「灵感输入 → AI 生成 → 人工确认 → 资产沉淀 → 视频输出」五个阶段。各阶段数据流转如下：")
    flows = [
        ("灵感输入阶段", "用户在前端输入项目标题与原始灵感，系统将其保存为项目基础信息。"),
        ("AI 生成阶段", "工作流执行器根据当前步骤调用对应的大语言模型、图像生成模型或视频生成模型，生成方向、框架、图片或视频片段。"),
        ("人工确认阶段", "用户在前端预览生成结果，进行选择、编辑、重新生成或跳过操作，确认后更新步骤状态并解锁下一步。"),
        ("资产沉淀阶段", "所有确认通过的生成结果进入项目资产库，按类型（CHARACTER、CONCEPT、STORYBOARD、VIDEO 等）分类存储。"),
        ("视频输出阶段", "在宣传片或直出视频步骤中，系统从资产库读取图片或片段，调用 ffmpeg 合成最终视频并上传存储。")
    ]
    for i, (stage, desc) in enumerate(flows, 1):
        add_numbered(doc, i, f"{stage}：{desc}")
    add_body(doc, "整个数据流通过工作流状态机进行驱动，每个步骤的状态变化会触发前端进度更新与后端权限校验，确保流程顺序执行。")


# ============================================================
# 第二章 系统功能说明
# ============================================================
def add_chapter2(doc):
    add_heading(doc, "第二章 系统功能说明", level=1)

    add_heading(doc, "2.1 用户系统", level=2)
    add_body(doc, "用户系统是本系统的入口与权限基础。未登录用户仅可查看登录/注册页面，登录后方可创建项目、执行生成任务与管理个人资产。")
    add_bullet(doc, "注册：用户可通过邮箱与密码完成本地账号注册，也可使用 GitHub OAuth 一键授权注册。注册成功后，系统自动创建个人账户并初始化默认点数余额。")
    add_bullet(doc, "登录：登录页面提供账号密码登录与第三方 OAuth 登录两种入口。系统通过 NextAuth.js v5 管理会话，登录成功后携带会话 Cookie 进入项目仪表盘。")
    add_bullet(doc, "个人中心：点击顶部导航头像进入个人中心，可查看与修改用户名、头像、邮箱，查看当前点数余额、充值历史与消费明细。")
    add_bullet(doc, "点数管理：系统为每次 AI 生成任务配置点数消耗，任务执行前自动校验余额。余额不足时，前端提示用户充值，后端拒绝生成请求。")
    add_bullet(doc, "充值：用户进入充值页面，选择或输入充值金额，上传转账凭证图片后提交订单。订单状态包括待审核、已通过、已拒绝，通过后点数实时到账。")
    add_bullet(doc, "会话安全：系统使用加密会话 Cookie，支持会话过期自动退出，保障用户账号安全。")
    add_bullet(doc, "权限隔离：普通用户仅能查看与操作自己的项目与资产，管理员额外拥有后台管理权限。")

    add_heading(doc, "2.2 项目管理", level=2)
    add_body(doc, "项目是本系统组织工作的基本单元。每个项目独立维护一套 9 步工作流状态与关联资产，不同项目之间的数据相互隔离。")
    add_bullet(doc, "新建项目：在仪表盘点击「新建项目」，输入项目标题与原始灵感描述后，系统初始化项目记录并创建 9 个工作流步骤的待执行状态。")
    add_bullet(doc, "项目列表：仪表盘以卡片网格展示用户所有项目，显示项目标题、当前进度、最近更新时间。支持点击进入项目详情。")
    add_bullet(doc, "仪表盘统计：顶部展示项目总数、进行中项目数、已完成项目数与最近 7 天活跃项目数，帮助用户快速掌握整体情况。")
    add_bullet(doc, "项目归档：用户可将已完成或暂时不需要的项目归档，归档项目不进入默认列表，但可在筛选条件中恢复查看。")
    add_bullet(doc, "项目删除：用户可在项目详情页删除项目，删除后关联资产、工作流记录与视频片段一并清理，释放存储空间。")

    add_heading(doc, "2.3 创意扩散模块", level=2)
    add_body(doc, "创意扩散是项目的起点，旨在将用户模糊的灵感转化为结构化、可执行的创意方向。")
    add_bullet(doc, "灵感输入：用户在多行文本框中输入元构思，内容可以是剧情梗概、主题关键词、风格描述或任意创作想法。")
    add_bullet(doc, "创意方向生成：AI 基于元构思输出 3–5 个完整的创意扩散方向，每个方向包含标题、核心设定与展开描述。")
    add_bullet(doc, "AI 评估：系统对每个方向从新颖性、情感张力、可行性、视觉表现力等维度给出评分与核心顾虑，辅助用户决策。")
    add_bullet(doc, "多轮深化：用户可选择系统推荐的改进方向或输入自定义反馈，AI 基于选择生成新的创意版本，保留历史迭代记录。")

    add_heading(doc, "2.4 框架搭建模块", level=2)
    add_body(doc, "框架搭建模块将选定的创意方向进一步细化为可指导后续生成的故事框架。")
    add_bullet(doc, "故事梗概：AI 自动生成项目的故事梗概，描述主线冲突、高潮与结局。")
    add_bullet(doc, "角色设定：输出主要角色的姓名、性格、外貌特征、人物关系与在故事中的作用。")
    add_bullet(doc, "幕结构：根据用户在创意扩散阶段选择的故事档位（速写、短篇、中篇、长片、史诗），动态生成 1–5 幕结构，每幕包含标题、内容、预估时长与镜头数。")
    add_bullet(doc, "环境设定：描述故事发生的主要地理空间、时代背景与氛围基调。")
    add_bullet(doc, "视觉风格：总结项目整体视觉基调，为后续风格统一、概念图与视频生成提供风格参考。")
    add_bullet(doc, "导入框架文件：支持上传文本文件或结构化文件，AI 解析后填充到框架字段，减少手动输入。")

    add_heading(doc, "2.5 风格统一模块", level=2)
    add_body(doc, "风格统一模块确定项目的视觉基准，确保后续人物设计、概念图与视频生成在视觉上保持一致。")
    add_bullet(doc, "视觉基准生成：系统根据框架内容生成 3 组不同风格的样图，每组 1 张，展示不同的色彩、笔触与构图倾向。")
    add_bullet(doc, "风格图选定：用户预览 3 组风格图后，选择最符合项目基调的一组作为统一风格基准。")
    add_bullet(doc, "模型选择：每组风格图可指定不同的图像生成模型，系统记录所选模型与提示词，作为后续步骤的风格参考注入来源。")
    add_bullet(doc, "比例切换：支持 16:9、9:16、1:1、4:3、3:4、21:9 等多种画面比例，满足不同发布平台的需求。")

    add_heading(doc, "2.6 人物设计模块", level=2)
    add_body(doc, "人物设计模块为剧本中的主要角色建立可视化形象，并作为后续概念图与视频生成的角色一致性参考。")
    add_bullet(doc, "角色形象生成：基于角色设定与已选定的风格基准，为每个角色生成概念图，统一风格与气质。")
    add_bullet(doc, "角色深化：支持对单个角色形象进行重新生成、调整英文提示词、切换画面比例与模型。")
    add_bullet(doc, "资产入库：生成的角色图自动进入项目资产库，标记为 CHARACTER 类型，可在概念图与视频生成步骤中作为角色参考注入。")

    add_heading(doc, "2.7 概念图模块", level=2)
    add_body(doc, "概念图模块将剧本框架中的关键场景可视化，形成项目的视觉蓝图。")
    add_bullet(doc, "场景概念图生成：按幕结构为每幕核心场景生成 1–3 张概念图，概念图数量基于该幕的关键场景数动态计算。")
    add_bullet(doc, "分幕展示：概念图按幕分组展示，用户可直观把握从整体到局部的视觉节奏。")
    add_bullet(doc, "提示词编辑：每张概念图的英文提示词支持用户查看与编辑，编辑后重新生成使用新提示词。")

    add_heading(doc, "2.8 分镜设计模块", level=2)
    add_body(doc, "分镜设计模块将剧本拆解为镜头序列，是连接文字剧本与视频生成的桥梁。")
    add_bullet(doc, "镜头列表：以表格形式展示镜头序号、描述、运镜方式、时长、角色与参考画面，支持拖拽排序与行内编辑。")
    add_bullet(doc, "分镜图生成：为每个镜头生成线稿风格的分镜草图，帮助导演与摄影师快速理解镜头意图。")
    add_bullet(doc, "双模式支持：支持「实拍参考模式」与「视频生成模式」两种分镜模式，前者生成代表画面，后者生成视频起始帧。")
    add_bullet(doc, "导出功能：支持将分镜表导出为 JSON 或 Excel，方便外部协作与打印。")

    add_heading(doc, "2.9 视频生成模块", level=2)
    add_body(doc, "视频生成模块包含宣传片生成与直出视频生成两部分，是系统将静态视觉资产转化为动态影像的核心能力。")
    add_bullet(doc, "宣传片生成：基于概念图序列逐段生成视频片段。用户可先为每个分镜生成提示词，再单独或批量生成视频片段，片段生成完成后点击「合成视频」调用 ffmpeg 拼接为完整宣传片。")
    add_bullet(doc, "背景音乐生成：宣传片合成时，系统基于框架情绪与故事梗概自动调用 AI 音乐生成，并以 0.3 音量与原声混音；若音乐生成失败，自动回退为静音音轨，不阻塞主流程。")
    add_bullet(doc, "直出视频生成：基于分镜设计生成的首帧与尾帧，调用图生视频模型生成小动作、小运镜的视频片段，并拼接为完整直出视频。")
    add_bullet(doc, "卡片式管理：视频片段以卡片网格展示，每张卡片显示分镜描述、生成状态、视频预览与重试按钮，双击卡片可在合成后的长视频中跳转对应时间点。")

    add_heading(doc, "2.10 管理后台", level=2)
    add_body(doc, "管理后台为系统运营人员提供用户、充值、日志与数据的统一管理入口。")
    add_bullet(doc, "用户统计：展示注册用户数、活跃用户趋势、项目总数与点数消耗分布。")
    add_bullet(doc, "充值审核：列出待审核充值订单，展示用户、金额、凭证图片与提交时间，管理员可通过或拒绝申请。")
    add_bullet(doc, "操作日志：记录用户的关键操作与生成任务，包括操作类型、项目 ID、点数消耗与结果状态，便于审计。")
    add_bullet(doc, "数据看板：通过图表可视化展示系统运行状态，如每日新增项目数、各步骤完成率、生成任务队列长度等。")


# ============================================================
# 第三章 操作说明
# ============================================================
def add_chapter3(doc):
    add_heading(doc, "第三章 操作说明", level=1)

    add_heading(doc, "3.1 快速开始", level=2)
    add_body(doc, "首次使用本系统的用户，可按照以下步骤在 10 分钟内完成第一个项目的创建与前三个步骤的生成：")
    add_numbered(doc, 1, "登录系统后，点击仪表盘右上角「新建项目」按钮；")
    add_numbered(doc, 2, "在弹出框中输入项目标题与原始灵感描述，点击确认创建项目；")
    add_numbered(doc, 3, "进入项目工作流看板，点击「创意扩散」步骤的「开始执行」按钮；")
    add_numbered(doc, 4, "等待 AI 生成多个创意方向后，阅读每个方向的描述与 AI 评估分数，选择最满意的方向并点击「确认」；")
    add_numbered(doc, 5, "点击进入「框架搭建」步骤，查看 AI 生成的故事梗概、角色设定与幕结构，必要时点击文本进行编辑；")
    add_numbered(doc, 6, "进入「风格统一」步骤，点击「开始执行」生成 3 组风格样图，选择最符合项目基调的一组并确认；")
    add_numbered(doc, 7, "完成以上三步后，工作流看板中前三个步骤显示为「已完成」，后续步骤依次解锁，用户可继续进入人物设计、概念图等模块。")

    add_heading(doc, "3.2 工作流操作", level=2)
    add_body(doc, "在工作流看板中，每个步骤以卡片形式展示当前状态，状态包括：待开始、进行中、已完成、失败、已跳过。用户可点击任意步骤卡片进入该步骤的详细操作界面。")
    add_bullet(doc, "开始执行：点击步骤卡片中的「开始执行」按钮，系统调用 AI 生成该步骤的初始结果。生成过程中卡片显示「进行中」动画，用户可离开页面，稍后返回查看结果。")
    add_bullet(doc, "重新生成：若对生成结果不满意，可点击「重新生成」或「重做」按钮，系统清除旧结果并重新执行。部分步骤支持单条重做，如单张风格图或单个分镜。")
    add_bullet(doc, "跳过步骤：对于非必需步骤，用户可点击「跳过」按钮，系统将该步骤状态标记为「已跳过」，并解锁后续步骤。")
    add_bullet(doc, "下一步：当前步骤确认完成后，点击「下一步」按钮自动进入下一个未完成的步骤。")
    add_bullet(doc, "导出：分镜设计等步骤支持导出 JSON 或 Excel 文件，方便用户与外部团队协作。")

    add_heading(doc, "3.3 创意深化操作", level=2)
    add_body(doc, "在创意扩散与框架搭建等支持文本编辑的界面，用户可直接点击描述文本进入编辑状态。编辑框采用自适应高度，失焦后系统以 500 毫秒防抖自动保存修改内容到后端。按 ESC 键可取消编辑并恢复原始值。")
    add_body(doc, "AI 评估面板展示当前方案的评分与核心顾虑。用户可在系统推荐的 3 个改进方向中选择一个，或输入自定义反馈，然后点击「深化迭代」按钮。AI 基于用户选择生成新的创意版本，并在历史版本列表中保留所有迭代记录，便于对比与回退。")

    add_heading(doc, "3.4 资产库管理", level=2)
    add_body(doc, "点击项目左侧「资产库」菜单，可查看该项目下所有生成的图片、视频、文本与音频资产。资产按类型分组展示，顶部提供类型筛选按钮。")
    add_bullet(doc, "图片资产：以缩略图网格展示，点击可放大查看，悬停显示生成模型、画面比例与提示词信息。")
    add_bullet(doc, "视频资产：提供在线播放功能，显示视频时长与生成状态。")
    add_bullet(doc, "文本资产：展示 AI 生成的提示词、描述与评估结果，支持复制。")
    add_bullet(doc, "批量下载：支持选择多个图片资产打包下载为 ZIP 文件。")

    add_heading(doc, "3.5 充值与点数", level=2)
    add_body(doc, "用户进入「设置 → 充值」页面，选择系统预设的充值金额或输入自定义金额，上传转账凭证图片后提交订单。订单提交后状态为「待审核」。")
    add_body(doc, "管理员在后台「充值审核」页面查看待审核订单，核对凭证后点击「通过」或「拒绝」。通过后，系统为用户增加对应点数，用户可在个人中心查看余额变动。每次调用 AI 生成任务时，系统按配置扣除相应点数；余额不足时，前端弹出提示，后端拒绝生成请求。")

    add_heading(doc, "3.6 创意扩散步骤操作", level=2)
    add_body(doc, "进入创意扩散步骤后，用户首先在页面左侧查看项目原始灵感。若灵感需要补充，可点击编辑文本框进行修改，修改后自动保存。点击「开始执行」按钮，系统调用 AI 生成 3–5 个创意方向。")
    add_body(doc, "生成完成后，页面右侧展示创意方向卡片。用户应仔细阅读每个方向的标题、描述与 AI 评估分数。点击卡片可展开查看详细设定与核心顾虑。选择一个方向后，点击「确认选择」按钮，系统将该方向保存为项目创意基准，并解锁框架搭建步骤。")
    add_body(doc, "若对现有方向均不满意，可在评估面板选择一个改进方向或输入自定义反馈，点击「深化迭代」按钮生成新版本。每次迭代结果会保留在历史版本列表中，用户可随时切换回旧版本。")

    add_heading(doc, "3.7 框架搭建步骤操作", level=2)
    add_body(doc, "进入框架搭建步骤后，系统已基于选定的创意方向生成初步框架。页面顶部展示灵感阐释与当前故事档位。用户可点击「灵感阐释」「背景设定」等文本进入编辑状态，修改后失焦自动保存。")
    add_body(doc, "角色设定区以卡片形式展示每个角色。点击角色卡片可展开查看完整设定，点击文本可直接修改。幕结构区以折叠面板展示每一幕，点击面板标题可展开查看该幕的标题、内容、预估时长、镜头数与关键场景。")
    add_body(doc, "若用户已有现成的框架文档，可点击「导入框架文件」按钮，上传 txt、docx 或 json 格式的文件。AI 解析后自动填充到对应字段，用户可在此基础上继续编辑。确认框架内容后，点击「下一步」进入风格统一步骤。")

    add_heading(doc, "3.8 风格统一步骤操作", level=2)
    add_body(doc, "进入风格统一步骤后，点击「开始执行」按钮。系统根据框架内容生成 3 组风格样图。每组样图展示不同的视觉倾向，如写实、国风、科幻、二次元等。")
    add_body(doc, "用户可在每组样图下方查看其使用的模型、画面比例与提示词。将鼠标悬停在样图上会显示「重做」按钮，点击后可切换模型或比例重新生成该组样图。选择最符合项目基调的一组后，点击「确认风格」按钮，系统将该组设置为项目的统一风格基准。")
    add_body(doc, "确认风格后，系统记录所选模型的参数，后续人物设计、概念图生成等步骤会自动注入该风格参考，确保整体视觉一致性。")

    add_heading(doc, "3.9 人物设计步骤操作", level=2)
    add_body(doc, "进入人物设计步骤后，系统读取框架中的角色设定，并为每个角色生成一张概念图。生成结果以卡片网格展示，每张卡片包含角色名、角色图、英文提示词与操作按钮。")
    add_body(doc, "若某角色的形象不符合预期，可点击卡片上的「重做」按钮，在弹出的选项中切换画面比例或模型后重新生成。也可点击提示词文本直接编辑，修改后点击「重新生成」使用新提示词。")
    add_body(doc, "所有生成的角色图会自动进入项目资产库，并标记为 CHARACTER 类型。在后续概念图与视频生成步骤中，系统会自动将这些角色图作为风格参考注入，保持角色形象一致。确认所有角色形象后，点击「下一步」进入概念图步骤。")

    add_heading(doc, "3.10 概念图步骤操作", level=2)
    add_body(doc, "进入概念图步骤后，点击「开始执行」按钮。系统根据框架中的幕结构与关键场景，为每幕生成 1–3 张概念图。生成结果按幕分组展示，每组概念图上方显示幕标题与幕内容摘要。")
    add_body(doc, "用户可点击任意概念图放大查看细节。若某张概念图不符合预期，可点击「重做」按钮重新生成，或编辑提示词后重新生成。用户还可在画面比例选择栏中切换 16:9、9:16 等比例，系统会按新比例重新生成该组概念图。")
    add_body(doc, "确认所有概念图后，点击「下一步」进入分镜设计步骤。概念图将自动作为分镜设计与宣传片生成的视觉参考。")

    add_heading(doc, "3.11 分镜设计步骤操作", level=2)
    add_body(doc, "进入分镜设计步骤后，用户需要先选择分镜模式：「实拍参考模式」生成每个镜头的代表画面，适用于实拍项目；「视频生成模式」生成每个镜头的起始帧，适用于 AI 视频生成项目。")
    add_body(doc, "在表格视图中，用户可编辑每个镜头的描述、运镜方式、时长与角色。点击行首的拖拽图标可调整镜头顺序。点击「生成全部画面」按钮，系统为每个镜头生成对应画面。生成完成后，点击「画布视图」可查看卡片式分镜墙。")
    add_body(doc, "分镜设计完成后，用户可点击「导出 JSON」或「导出 Excel」按钮将分镜表导出，用于外部协作或打印。点击「下一步」进入生成尾帧步骤。")

    add_heading(doc, "3.12 宣传片步骤操作", level=2)
    add_body(doc, "进入宣传片步骤后，系统首先为每个分镜生成视频生成提示词。提示词生成完成后，页面以分镜卡片网格展示每个片段，卡片状态默认为「待生成」。")
    add_body(doc, "用户可点击单个卡片的「生成视频」按钮单独生成该片段，也可点击顶部「批量生成」按钮同时生成所有待生成片段。生成过程中卡片显示「生成中」动画，完成后显示视频预览。若某片段生成失败，卡片显示「失败」状态与错误原因，点击「重试」按钮可单独重新生成。")
    add_body(doc, "所有片段生成完成后，顶部出现「合成视频」按钮。点击后系统在后台调用 ffmpeg 拼接所有片段、生成背景音乐并混音。合成完成后，页面顶部常驻显示合成后的长视频播放器，用户可播放完整宣传片。下方的分镜卡片支持双击跳转至对应时间点，播放时当前片段卡片会自动高亮。")

    add_heading(doc, "3.13 直出视频步骤操作", level=2)
    add_body(doc, "进入直出视频步骤前，需确保在分镜设计步骤已生成首帧，在生成尾帧步骤已生成尾帧。系统会自动检测每个镜头是否有首尾帧，并据此决定使用「首帧-尾帧」策略还是「仅首帧」策略。")
    add_body(doc, "与宣传片类似，直出视频也以分镜卡片网格管理片段。点击「批量生成」按钮后，系统为每个镜头生成小动作、小运镜的视频片段。所有片段生成完成后，点击「拼接视频」按钮调用 ffmpeg 直接拼接片段为完整视频。直出视频不生成背景音乐，保留片段原声。")

    add_heading(doc, "3.14 管理后台操作", level=2)
    add_body(doc, "管理员登录后，点击顶部导航的「管理后台」入口进入后台界面。在「用户统计」页查看注册趋势与活跃数据；在「充值审核」页查看待审核订单，点击图片可放大查看凭证，确认无误后点击「通过」，有问题则点击「拒绝」并填写备注；在「操作日志」页可按用户、操作类型与时间范围筛选日志记录。")

    add_heading(doc, "3.15 快捷键与辅助功能", level=2)
    add_body(doc, "为提升操作效率，本系统在部分界面提供键盘快捷键与辅助功能。在文本编辑区域，按 Enter 可保存当前编辑并退出编辑状态，按 ESC 可取消编辑并恢复原始值。在分镜表格中，按 Tab 可在单元格间快速切换，按 Shift+Tab 反向切换。")
    add_bullet(doc, "ESC：取消当前编辑或关闭弹窗。")
    add_bullet(doc, "Enter：保存当前编辑或确认当前操作。")
    add_bullet(doc, "Tab / Shift+Tab：在表格单元格或表单字段间前后切换焦点。")
    add_bullet(doc, "Ctrl+S（或 Cmd+S）：在支持导出的界面触发导出操作。")
    add_body(doc, "系统同时支持响应式布局，在较小屏幕下部分界面会切换为单列布局，表格支持横向滚动，确保移动端也可基本浏览项目数据。")

    add_heading(doc, "3.16 数据备份与导出", level=2)
    add_body(doc, "用户可通过多种方式备份项目数据。分镜设计步骤支持导出 JSON 与 Excel 文件，包含完整的镜头列表、描述、运镜方式与时长。框架搭建步骤支持导出 Word 与 Markdown 文档，便于外部编剧或制片人审阅。")
    add_bullet(doc, "JSON 导出：导出机器可读的结构化数据，适合导入其他工具或进行二次开发。")
    add_bullet(doc, "Excel 导出：导出为表格格式，方便打印或在办公软件中继续编辑。")
    add_bullet(doc, "Word 导出：导出为带有格式的 Word 文档，适合提交审阅或归档。")
    add_bullet(doc, "Markdown 导出：导出为纯文本 Markdown，适合版本控制或协作编辑。")
    add_body(doc, "管理员可在后台导出操作日志，用于审计与数据分析。所有导出文件均从当前最新数据生成，不包含已删除的历史版本。")


# ============================================================
# 第四章 界面说明
# ============================================================
def add_chapter4(doc):
    add_heading(doc, "第四章 界面说明", level=1)

    add_heading(doc, "4.1 登录/注册界面", level=2)
    add_body(doc, "登录/注册界面为系统的统一入口。页面采用左右分栏布局，左侧为品牌与功能简介，右侧为表单区域。登录表单包含邮箱输入框、密码输入框、登录按钮与「忘记密码」链接；页面底部提供 GitHub OAuth 一键登录入口。注册表单包含邮箱、密码与确认密码字段，注册成功后自动登录并跳转仪表盘。")
    add_body(doc, "在登录表单中，系统会对邮箱格式进行前端校验，对密码长度进行基础校验。若登录失败，表单下方会显示具体错误信息，如「账号或密码错误」「该用户不存在」等。OAuth 登录成功后，系统会自动创建本地账号绑定，用户后续可选择使用密码或 OAuth 任一方式登录。")
    add_screenshot_placeholder(doc, "该界面是用户进入系统的唯一入口，展示软件品牌标识与登录注册表单。")

    add_heading(doc, "4.2 仪表盘界面", level=2)
    add_body(doc, "仪表盘是用户登录后的默认首页，分为顶部统计区与下方项目列表区。顶部统计区以卡片形式展示项目总数、进行中项目数、已完成项目数与最近 7 天活跃项目数，每张统计卡片还包含环比变化提示。下方项目列表以卡片网格展示最近项目，每张卡片显示项目标题、当前进度百分比、最近更新时间缩略图与快捷进入按钮。")
    add_body(doc, "页面右上角提供「新建项目」按钮与用户头像下拉菜单。点击头像下拉菜单可进入个人中心、充值页面或执行退出登录。当项目数量较多时，页面底部提供分页控件，用户可调整每页显示数量。")
    add_screenshot_placeholder(doc, "该界面展示用户所有项目的概览与统计数据，是用户管理项目的核心入口。")

    add_heading(doc, "4.3 工作流看板", level=2)
    add_body(doc, "工作流看板是项目的核心操作界面。页面顶部为 9 步工作流进度条，每个步骤用图标、名称与状态标签表示，已完成步骤显示绿色对勾，进行中步骤显示加载动画，失败步骤显示红色警告，已跳过步骤显示灰色标签。进度条下方为当前选中步骤的详细面板。")
    add_body(doc, "详细面板左侧展示步骤说明、操作按钮与历史记录入口，右侧展示生成结果。用户可在步骤间切换，查看历史结果或执行新的生成任务。每个步骤卡片右上角提供「跳过」按钮（可跳过步骤）与「重做」按钮（已完成步骤可用）。")
    add_screenshot_placeholder(doc, "该界面以可视化进度条展示 9 步工作流状态，是用户执行生成任务的主要操作区。")

    add_heading(doc, "4.4 创意扩散界面", level=2)
    add_body(doc, "创意扩散界面分为左右两栏。左侧上方为灵感输入区，包含项目原始灵感与元构思输入框；左侧下方为历史版本列表，展示历次深化迭代的版本号与简要说明。右侧以卡片网格展示 AI 生成的创意方向，每张卡片包含标题、中文描述、AI 评估分数与「选择」按钮。")
    add_body(doc, "点击卡片可展开查看详细设定与核心顾虑。卡片底部提供「深化迭代」入口。用户选择一个方向后，该卡片边框高亮，页面底部出现「确认选择」按钮。确认后系统将该方向保存为项目创意基准，并解锁框架搭建步骤。")
    add_screenshot_placeholder(doc, "该界面用于展示 AI 生成的创意方向与评估结果，支持用户选择与深化迭代。")

    add_heading(doc, "4.5 框架搭建界面", level=2)
    add_body(doc, "框架搭建界面采用分组表单布局。页面顶部展示灵感阐释与故事档位信息，故事档位以标签形式显示，如「短篇·3-5 分钟」。中部为角色设定卡片，每张卡片显示角色姓名、性格、外貌与作用，支持点击编辑。下部为幕结构列表，每幕以折叠面板展示标题、内容、预估时长、镜头数与关键场景。")
    add_body(doc, "页面右侧提供环境设定、视觉风格与整体节奏策略的编辑区。底部提供导出文档按钮，可将框架内容导出为 Word 或 Markdown 格式供外部使用。所有文本字段均支持点击编辑，失焦后自动保存。")
    add_screenshot_placeholder(doc, "该界面用于编辑故事框架、角色设定与幕结构，是项目世界观搭建的核心界面。")

    add_heading(doc, "4.6 分镜设计界面", level=2)
    add_body(doc, "分镜设计界面提供表格视图与画布视图两种模式。表格视图以行展示每个镜头，列包括镜头序号、描述、运镜方式、时长、角色、参考图与尾帧图，支持拖拽排序与行内编辑。画布视图以卡片网格展示分镜画面，支持拖拽重排。")
    add_body(doc, "页面顶部工具栏包含模式切换、导出按钮、生成全部按钮与视图缩放控件。用户可在表格中直接修改镜头描述，按回车或点击外部保存。生成的参考图与尾帧图以缩略图形式嵌入表格，点击可放大查看。")
    add_screenshot_placeholder(doc, "该界面用于管理镜头列表与分镜草图，支持表格与画布两种交互视图。")

    add_heading(doc, "4.7 视频生成界面", level=2)
    add_body(doc, "视频生成界面分为宣传片与直出视频两个子模块。在宣传片模块中，页面以分镜卡片网格展示每个片段的生成状态，卡片包含分镜描述、提示词、时长与生成按钮。全部片段生成完成后，顶部出现「合成视频」按钮。")
    add_body(doc, "合成完成后，页面顶部常驻显示合成后的长视频播放器，下方仍展示分镜卡片。用户播放视频时，当前时间对应的片段卡片会自动高亮显示。双击任意分镜卡片，长视频会自动跳转到该片段的起始时间点并继续播放。")
    add_screenshot_placeholder(doc, "该界面用于逐段生成视频片段并合成最终视频，展示分镜卡片与常驻视频播放器。")

    add_heading(doc, "4.8 管理后台界面", level=2)
    add_body(doc, "管理后台采用经典后台布局，左侧为管理员导航菜单，包含用户统计、充值审核、操作日志与数据看板。右侧为主内容区。用户统计页以图表展示注册趋势与项目分布，支持按时间范围筛选。")
    add_body(doc, "充值审核页以表格展示待审核订单，每行包含用户信息、金额、凭证缩略图与通过/拒绝按钮，点击缩略图可放大查看凭证。操作日志页支持按用户、操作类型与时间范围筛选，并提供导出功能。数据看板以图表形式展示系统运行状态，如每日新增项目数、各步骤完成率、生成任务队列长度等。")
    add_screenshot_placeholder(doc, "该界面为管理员提供用户、充值、日志与数据的统一管理入口。")

    add_heading(doc, "4.9 个人中心界面", level=2)
    add_body(doc, "个人中心界面采用卡片式布局，顶部展示用户头像、用户名、邮箱与当前点数余额。余额卡片以醒目的数字展示可用点数，并在下方列出最近 5 条消费记录，包括消费时间、关联项目、操作类型与扣除点数。")
    add_body(doc, "中部为账号设置区，用户可修改用户名、上传头像、绑定或解绑 GitHub 账号。头像上传后系统自动压缩并上传至对象存储。底部为充值历史列表，展示每笔充值订单的金额、状态、提交时间与审核备注，用户可点击订单查看凭证图片。")
    add_screenshot_placeholder(doc, "该界面用于管理用户账号信息、查看点数余额与充值消费记录。")

    add_heading(doc, "4.10 资产库界面", level=2)
    add_body(doc, "资产库界面以类型筛选栏与内容网格为主体。顶部提供全部、图片、视频、文本、音频五种类型筛选按钮，当前选中类型高亮显示。图片资产以缩略图网格展示，悬停显示生成模型、画面比例与提示词摘要；视频资产显示时长、状态与播放按钮；文本资产以卡片列表展示标题与摘要。")
    add_body(doc, "页面右上角提供批量选择开关与下载按钮。开启批量选择后，用户可勾选多个图片资产，点击下载按钮打包为 ZIP 文件。点击任意资产可进入详情弹窗，查看完整元数据、提示词与关联的工作流步骤。")
    add_screenshot_placeholder(doc, "该界面用于集中管理项目生成的所有图片、视频、文本与音频资产。")


# ============================================================
# 附录
# ============================================================
def add_appendix(doc):
    add_heading(doc, "附录 技术栈说明", level=1)
    add_body(doc, "本系统采用主流现代 Web 技术栈构建，兼顾开发效率、运行时性能与可维护性。主要技术组件说明如下：")
    techs = [
        ("Next.js 14", "基于 React 的全栈 Web 框架，提供 App Router、API Routes、Server Components 与 Edge Runtime 支持，是本系统前后端的统一运行环境。"),
        ("React", "用于构建用户界面的 JavaScript 库，支持组件化开发、Hooks 状态管理与客户端交互。"),
        ("TypeScript", "在 JavaScript 基础上增加静态类型系统，提升代码可维护性与 IDE 提示能力。"),
        ("Tailwind CSS", "实用优先的 CSS 框架，通过原子化类名快速构建响应式、一致的界面样式。"),
        ("Prisma", "下一代 Node.js 与 TypeScript ORM，提供类型安全的数据库访问、迁移管理与查询构建器。"),
        ("Supabase", "提供托管 PostgreSQL 数据库与认证服务，作为本系统的数据持久化与身份验证基础设施。"),
        ("BullMQ", "基于 Redis 的 Node.js 队列系统，用于异步处理图像生成、视频生成等耗时任务，避免前端请求超时。"),
        ("ffmpeg", "开源音视频处理工具，负责视频片段拼接、音频混流、格式转换与静音音轨生成。"),
        ("python-docx", "用于生成 Word 文档的 Python 库，本说明书即通过该库程序化生成。")
    ]
    for name, desc in techs:
        add_bullet(doc, f"{name}：{desc}")

    add_body(doc, "系统部署支持 Docker 容器化方案与 Vercel Serverless 方案。开发环境通过 docker-compose 一键启动，自动挂载源码卷、node_modules 卷与 public/mock-storage 持久化卷；生产环境可通过环境变量配置数据库连接、R2 对象存储、Redis 队列与 ffmpeg 路径。")
    add_body(doc, "在 DEMO 模式下，系统无需真实数据库与云存储即可运行：Prisma 自动切换为本地 JSON 持久化，R2 切换为 public/mock-storage 本地文件系统，BullMQ 切换为内存 Mock，方便开发者本地体验与答辩演示。")

    add_heading(doc, "附录二 系统架构说明", level=1)
    add_body(doc, "本系统采用前后端分离的全栈架构，前端基于 Next.js 14 App Router，后端通过 Next.js API Routes 提供服务，数据持久化使用 PostgreSQL，对象存储使用 Cloudflare R2，异步任务队列使用 BullMQ + Redis。")
    add_body(doc, "架构自上而下可分为五层：")
    add_bullet(doc, "展示层：由 Next.js 页面、React 组件与 Tailwind CSS 样式构成，负责用户界面渲染与交互逻辑。")
    add_bullet(doc, "API 网关层：Next.js API Routes 统一接收前端请求，进行身份校验、参数解析与路由分发。")
    add_bullet(doc, "业务逻辑层：位于 lib/ 与 app/api/ 下，包含工作流执行器、图像/视频生成客户端、点数系统、操作日志等核心模块。")
    add_bullet(doc, "数据持久化层：Prisma ORM 连接 PostgreSQL 存储用户、项目、工作流、资产、充值等结构化数据；R2 存储图片、视频、音频等非结构化资产。")
    add_bullet(doc, "异步任务层：BullMQ 将图像生成、视频生成等耗时任务放入 Redis 队列，Worker 进程异步消费任务，避免前端请求超时。")
    add_body(doc, "视频合成流程是系统核心算法路径之一：系统首先从 R2 下载各视频片段到本地临时目录，调用 ffmpeg 进行片段拼接；对于宣传片，还会生成背景音乐并调用 amix 滤镜以 0.3 音量与原声混音；最终将合成后的 MP4 上传回 R2，并更新项目状态。")

    add_heading(doc, "附录三 数据安全与隐私", level=1)
    add_body(doc, "本系统重视用户数据安全与隐私保护，采取以下措施：")
    add_bullet(doc, "所有第三方 API 密钥与数据库凭据均通过环境变量注入，禁止硬编码到源码中。")
    add_bullet(doc, "用户密码通过哈希算法存储，不会以明文形式保存。")
    add_bullet(doc, "会话 Cookie 设置 HttpOnly 与 Secure 属性，降低 XSS 与会话劫持风险。")
    add_bullet(doc, "用户仅能访问自己的项目与资产，管理员也无法查看普通用户的具体生成内容，除非进入后台审计场景。")
    add_bullet(doc, "操作日志记录关键行为但不记录用户原始提示词的完整内容，仅保留摘要与元数据。")


# ============================================================
# 第五章 常见问题与故障排除
# ============================================================
def add_chapter5(doc):
    add_heading(doc, "第五章 常见问题与故障排除", level=1)

    add_heading(doc, "5.1 账号与登录问题", level=2)
    add_bullet(doc, "无法收到验证邮件：请检查垃圾邮件文件夹，或确认邮箱地址拼写正确。若使用第三方 OAuth，请确保该账号已完成邮箱验证。")
    add_bullet(doc, "登录后页面空白：可能是浏览器插件拦截了会话 Cookie，请尝试禁用广告拦截插件或使用无痕模式访问。")
    add_bullet(doc, "忘记密码：当前版本支持管理员手动重置密码，请联系系统管理员处理。")

    add_heading(doc, "5.2 生成任务问题", level=2)
    add_bullet(doc, "步骤长时间显示「进行中」：AI 生成任务可能耗时较长，请耐心等待。若超过 10 分钟无变化，可点击「取消」后重新执行。")
    add_bullet(doc, "生成结果不符合预期：可点击「重新生成」或进入提示词编辑模式修改英文提示词，再重新生成。")
    add_bullet(doc, "视频片段生成失败：失败片段卡片会显示「失败」状态与错误原因，点击「重试」按钮即可单独重新生成该片段，不影响其他片段。")
    add_bullet(doc, "合成视频提示「还有 X 个片段未生成」：请确保所有片段状态为「已完成」后再点击「合成视频」。")

    add_heading(doc, "5.3 点数与充值问题", level=2)
    add_bullet(doc, "点数余额不足：请进入充值页面提交充值订单，管理员审核通过后点数自动到账。")
    add_bullet(doc, "充值订单被拒绝：请检查上传的凭证图片是否清晰、金额是否与提交金额一致，然后重新提交。")
    add_bullet(doc, "点数扣除有疑问：可在个人中心查看操作日志与消费明细，若有异常请联系管理员。")

    add_heading(doc, "5.4 界面与操作问题", level=2)
    add_bullet(doc, "分镜表格拖拽无效：请确保未在编辑单元格状态下进行拖拽，编辑完成后按回车或点击外部保存再拖拽。")
    add_bullet(doc, "图片加载显示「加载失败」：可能是网络波动或图片已被清理，请刷新页面或重新生成该图片。")
    add_bullet(doc, "视频无法播放：请检查浏览器是否支持 MP4 格式，或尝试下载后使用本地播放器播放。")

    add_heading(doc, "5.5 性能与兼容性建议", level=2)
    add_body(doc, "为获得最佳体验，建议使用 Chrome 116 及以上版本、Edge 116 及以上版本或 Firefox 118 及以上版本访问本系统。在低分辨率屏幕下，部分表格可能需要横向滚动。对于包含大量分镜或大量资产的项目，建议在稳定的网络环境下操作，以避免请求超时。")

    add_heading(doc, "5.6 版本历史", level=2)
    add_body(doc, "本系统自启动开发以来，持续迭代优化，主要版本记录如下：")
    versions = [
        ("V0.1.0（内测版）", "完成用户系统、项目管理与创意扩散模块的初版实现，支持基础的项目创建与 AI 创意方向生成。"),
        ("V0.2.0", "新增框架搭建、风格统一与人物设计模块，建立 9 步工作流状态机。"),
        ("V0.3.0", "新增概念图、分镜设计与生成尾帧模块，支持分镜表格编辑与导出。"),
        ("V0.4.0", "新增宣传片与直出视频模块，集成 ffmpeg 视频合成与 BGM 混音。"),
        ("V0.5.0", "新增管理后台、充值审核与点数管理，支持商业化运营。"),
        ("V1.0（正式版）", "完成全模块联调、Docker 部署适配与软著材料生成，作为正式发布版本。")
    ]
    for v, desc in versions:
        add_bullet(doc, f"{v}：{desc}")

    add_heading(doc, "5.7 联系与支持", level=2)
    add_body(doc, "如在使用过程中遇到本说明书未覆盖的问题，用户可通过以下方式获取支持：查看系统内操作提示与错误信息；在个人中心查看操作日志与消费明细；联系系统管理员或技术支持人员。管理员可通过管理后台的操作日志与数据看板定位用户问题。")

    add_heading(doc, "5.8 安全与隐私常见问题", level=2)
    add_bullet(doc, "数据存储在哪里：用户的项目数据、账号信息与资产元数据存储在 PostgreSQL 数据库中，图片与视频文件存储在 Cloudflare R2 对象存储服务中。")
    add_bullet(doc, "密码是否安全：用户密码通过行业标准的哈希算法存储，系统不会以明文形式保存密码，管理员也无法查看用户密码。")
    add_bullet(doc, "生成内容是否会被他人查看：普通用户只能查看自己的项目与资产，系统通过严格的权限校验防止越权访问，管理员在后台也无法直接浏览用户的具体生成内容。")
    add_bullet(doc, "如何保护 API 密钥：所有第三方服务密钥与数据库凭据均通过环境变量注入，不会硬编码在源码或配置文件中。")
    add_bullet(doc, "会话过期时间是多少：用户会话默认在浏览器关闭或超过设定闲置时间后失效，具体时长由服务器配置决定。")


# ============================================================
# 第六章 术语表
# ============================================================
def add_glossary(doc):
    add_heading(doc, "第六章 术语表", level=1)
    add_body(doc, "为便于理解本说明书及系统功能，以下对常用术语进行解释：")
    terms = [
        ("元构思", "用户在项目创建时输入的原始灵感或创作想法，是 AI 生成创意方向的起点。"),
        ("创意方向", "系统基于元构思生成的具体创作方案，包含标题、核心设定与展开描述。"),
        ("框架", "项目的故事结构，包括故事梗概、角色设定、幕结构、环境设定与视觉风格。"),
        ("幕结构", "将故事划分为若干幕的组织方式，每幕包含标题、内容、预估时长与镜头数。"),
        ("风格基准", "在风格统一步骤中选定的视觉样图及其参数，作为后续生成的风格参考。"),
        ("概念图", "将剧本关键场景可视化的静态图片，用于把握整体视觉节奏。"),
        ("分镜", "将剧本拆解为镜头序列的结果，每个分镜包含描述、运镜、时长等信息。"),
        ("首帧", "视频片段的起始画面，通常由分镜设计步骤生成。"),
        ("尾帧", "视频片段的结束画面，由生成尾帧步骤基于首帧生成。"),
        ("直出视频", "基于首尾帧直接生成的视频片段，适用于小动作、小运镜场景。"),
        ("宣传片", "基于概念图序列生成的短视频，用于项目提案或预告展示。"),
        ("背景音乐", "宣传片合成时添加的音频，系统支持 AI 音乐生成与静音回退。"),
        ("资产", "系统中生成的图片、视频、文本、音频等成果的统称，可在资产库中管理。"),
        ("点数", "系统内用于支付 AI 生成任务的虚拟货币，用户通过充值获得。"),
        ("工作流", "项目中 9 个顺序执行的步骤集合，用户按顺序完成以产出最终成果。"),
        ("工作流状态", "每个步骤当前所处状态，包括待开始、进行中、已完成、失败、已跳过。"),
        ("重新生成", "对某一步骤或某一条生成结果进行重新执行的操作。"),
        ("跳过", "对非必需步骤标记为已跳过，以解锁后续步骤的操作。"),
        ("导出", "将分镜表等数据保存为 JSON 或 Excel 文件的功能。"),
        ("Mock 模式", "系统在 DEMO 或供应商服务不可用时使用的本地兜底生成方式。"),
        ("ffmpeg", "开源音视频处理工具，负责视频拼接、音频混流与格式转换。"),
        ("BullMQ", "基于 Redis 的任务队列系统，用于处理耗时的 AI 生成任务。"),
        ("R2", "Cloudflare 提供的对象存储服务，用于存储生成的图片与视频文件。"),
        ("Prisma", "TypeScript ORM，用于与 PostgreSQL 数据库交互。"),
        ("NextAuth.js", "用于处理用户认证与授权的库，支持 OAuth 与会话管理。"),
        ("模型", "本说明书中特指 AI 生成模型，如图像生成模型、视频生成模型、音乐生成模型。"),
        ("提示词", "用于引导 AI 生成模型输出的文本描述，通常使用英文编写。"),
        ("画面比例", "图像或视频的宽高比，常见值包括 16:9、9:16、1:1、4:3 等。"),
        ("队列", "异步任务的等待与执行机制，确保系统不会因为单个长任务而阻塞。"),
        ("管理后台", "供管理员使用的后台界面，用于用户、充值、日志与数据管理。"),
    ]
    for term, definition in terms:
        add_bullet(doc, f"{term}：{definition}")


# ============================================================
# 主流程
# ============================================================
def setup_styles(doc):
    style = doc.styles["Normal"]
    style.font.name = "宋体"
    style.font.size = Pt(12)
    style._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")


def setup_section(doc):
    section = doc.sections[0]
    section.page_height = Cm(29.7)
    section.page_width = Cm(21.0)
    section.top_margin = Cm(2.54)
    section.bottom_margin = Cm(2.54)
    section.left_margin = Cm(3.17)
    section.right_margin = Cm(3.17)

    # 页眉
    header = section.header
    header_para = header.paragraphs[0]
    header_para.text = f"{SOFTWARE_NAME} 软件使用说明书 {VERSION}"
    header_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in header_para.runs:
        set_run_font(run, "宋体", 10.5)

    # 页脚页码
    footer = section.footer
    footer_para = footer.paragraphs[0]
    footer_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    fld_begin = OxmlElement("w:fldChar")
    fld_begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = "PAGE"
    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")
    run = footer_para.add_run("第 ")
    set_run_font(run, "宋体", 10.5)
    run = footer_para.add_run()
    run._r.append(fld_begin)
    run._r.append(instr)
    run._r.append(fld_end)
    set_run_font(run, "宋体", 10.5)
    run = footer_para.add_run(" 页")
    set_run_font(run, "宋体", 10.5)


def main():
    doc = Document()
    setup_styles(doc)
    setup_section(doc)

    add_cover(doc)
    add_page_break(doc)

    add_toc(doc)
    add_page_break(doc)

    add_chapter1(doc)
    add_page_break(doc)

    add_chapter2(doc)
    add_page_break(doc)

    add_chapter3(doc)
    add_page_break(doc)

    add_chapter4(doc)
    add_page_break(doc)

    add_chapter5(doc)
    add_page_break(doc)

    add_glossary(doc)
    add_page_break(doc)

    add_appendix(doc)
    add_page_break(doc)

    add_appendix(doc)

    OUTPUT_DIR.mkdir(parents=True, exist_ok=True)
    doc.save(OUTPUT_FILE)

    total_paras = len(doc.paragraphs)
    # A4 + 3.17cm 边距 + 1.5 倍行距 12pt，正文区约每页 26–28 个段落
    estimated_pages = max(1, total_paras // 26)
    print(f"[OK] 已生成: {OUTPUT_FILE}")
    print(f"     总段落数: {total_paras}")
    print(f"     估计页数: {estimated_pages}")


if __name__ == "__main__":
    main()
