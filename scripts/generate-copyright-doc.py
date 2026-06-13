#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
生成软著申请材料：源代码前/后 30 页 Word 文档。

用法：
  python scripts/generate-copyright-doc.py

输出：
  /mnt/agents/output/软著材料/源代码前30页.docx
  /mnt/agents/output/软著材料/源代码后30页.docx
"""

import re
import os
import sys
from pathlib import Path
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.section import WD_ORIENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# ============================================================
# 配置
# ============================================================
PROJECT_ROOT = Path(__file__).resolve().parent.parent


def resolve_posix_path(posix_path: str) -> Path:
    """在 Windows Git Bash 中通过 cygpath 转换 POSIX 路径；其他环境直接使用。"""
    import subprocess

    try:
        result = subprocess.run(
            ["cygpath", "-w", posix_path],
            capture_output=True,
            text=True,
            encoding="utf-8",
            errors="replace",
            check=True,
            timeout=5,
        )
        resolved = Path(result.stdout.strip())
        print(f"[INFO] 通过 cygpath 解析输出目录: {resolved}")
        return resolved
    except Exception as e:
        print(f"[INFO] 未使用 cygpath，直接使用路径: {posix_path} ({e})")
        return Path(posix_path)


OUTPUT_DIR = resolve_posix_path("/mnt/agents/output/软著材料")

SOFTWARE_NAME = "AI影视全流程工作流系统"
SOFTWARE_VERSION = "V1.0"
LINES_PER_PAGE = 50
MAX_LINES = 30 * LINES_PER_PAGE  # 1500 行

# 前 30 页固定文件清单（按顺序读取，自然截断至 30 页/1500 行）
FILE_ORDER_FRONT = [
    "package.json",
    "prisma/schema.prisma",
    "lib/workflow.ts",
    "lib/workflow-state.ts",
    "app/api/projects/route.ts",
    "app/api/projects/[id]/route.ts",
    "app/api/projects/[id]/steps/ideation/route.ts",
    "app/api/projects/[id]/steps/framework/route.ts",
    "app/api/projects/[id]/steps/style/route.ts",
    "app/api/projects/[id]/steps/character/route.ts",
    "app/api/projects/[id]/steps/concept/route.ts",
    "app/api/projects/[id]/steps/storyboard/route.ts",
    "app/api/projects/[id]/steps/trailer/route.ts",
    "app/api/projects/[id]/steps/video-direct/route.ts",
    "app/api/projects/[id]/steps/keyframes/route.ts",
    "app/api/projects/[id]/steps/keyframes/generate-last/route.ts",
]

# 后 30 页固定文件清单（接续前 30 页，排除已纳入文件，自然截断至 30 页/1500 行）
FILE_ORDER_BACK = [
    "lib/video-utils.ts",
    "lib/video-segment-utils.ts",
    "lib/bgm-generator.ts",
    "lib/workflow-executor.ts",
    "lib/models-config.ts",
    "app/(dashboard)/project/[id]/workflow/page.tsx",
    "app/(dashboard)/project/[id]/workflow/_components/TopStepper.tsx",
    "app/(dashboard)/project/[id]/storyboard/page.tsx",
    "app/(dashboard)/project/[id]/storyboard/_components/StoryboardTable.tsx",
    "app/(dashboard)/project/[id]/storyboard/_components/StoryboardCanvas.tsx",
    "components/workflow/IdeaAnchor.tsx",
    "components/workflow/StepBadge.tsx",
    "app/admin/analytics/page.tsx",
    "app/admin/users/page.tsx",
    "app/admin/recharges/page.tsx",
    "app/api/admin/users/[id]/points/route.ts",
]


# ============================================================
# 脱敏
# ============================================================
SENSITIVE_PATTERNS = [
    # API Key / Token / Secret
    (re.compile(r"\b(sk-[a-zA-Z0-9_\-]{20,})\b", re.IGNORECASE), "***API_KEY***"),
    (re.compile(r"\b([a-zA-Z0-9_\-]*api[_\-]?key[a-zA-Z0-9_\-]*\s*[:=]\s*[\"']?)([a-zA-Z0-9_\-]{8,})([\"']?)", re.IGNORECASE), r"\1***API_KEY***\3"),
    (re.compile(r"\b(Bearer\s+)([a-zA-Z0-9_\-\.]+)\b", re.IGNORECASE), r"\1***TOKEN***"),
    (re.compile(r"\b([a-zA-Z0-9_\-]*secret[a-zA-Z0-9_\-]*\s*[:=]\s*[\"']?)([a-zA-Z0-9_\-/+=]{8,})([\"']?)", re.IGNORECASE), r"\1***SECRET***\3"),
    # 数据库连接字符串
    (re.compile(r"\b(postgresql://[^\"'`\s]+)\b", re.IGNORECASE), "postgresql://***DB_URL***"),
    (re.compile(r"\b(mongodb(\+srv)?://[^\"'`\s]+)\b", re.IGNORECASE), "mongodb://***DB_URL***"),
    (re.compile(r"\b(mysql://[^\"'`\s]+)\b", re.IGNORECASE), "mysql://***DB_URL***"),
    # 邮箱
    (re.compile(r"\b[A-Za-z0-9._%+-]+@[A-Za-z0-9.-]+\.[A-Za-z]{2,}\b"), "***EMAIL***"),
    # 手机号（中国大陆）
    (re.compile(r"\b1[3-9]\d{9}\b"), "***PHONE***"),
    # 身份证号
    (re.compile(r"\b\d{17}[\dXx]|\d{15}\b"), "***ID_CARD***"),
    # URL 中的用户名密码
    (re.compile(r"(https?://)[^:@/\s]+:[^@/\s]+@", re.IGNORECASE), r"\1***CREDS***@"),
]


def sanitize_line(line: str) -> str:
    """对单行代码进行脱敏处理。"""
    for pattern, replacement in SENSITIVE_PATTERNS:
        line = pattern.sub(replacement, line)
    return line


# ============================================================
# 文件读取与截断
# ============================================================
def read_file_lines(rel_path: str) -> list[str]:
    """读取单个文件并脱敏，返回行列表。"""
    file_path = PROJECT_ROOT / rel_path
    if not file_path.exists():
        print(f"[WARN] 文件不存在，跳过: {rel_path}")
        return []
    try:
        with open(file_path, "r", encoding="utf-8") as f:
            raw_lines = f.read().splitlines()
    except UnicodeDecodeError:
        with open(file_path, "r", encoding="utf-8", errors="replace") as f:
            raw_lines = f.read().splitlines()

    lines = [sanitize_line(line) for line in raw_lines]
    # 文件分隔行
    separator = f"/* ===== {rel_path} ===== */"
    return [separator] + lines


def collect_content(file_order: list[str], max_lines_per_file: int | None = None) -> tuple[list[str], list[str]]:
    """
    按顺序收集文件内容，直到凑满 1500 行。
    返回 (lines, included_files)。
    """
    all_lines: list[str] = []
    included_files: list[str] = []

    for rel_path in file_order:
        if len(all_lines) >= MAX_LINES:
            break
        file_lines = read_file_lines(rel_path)
        if not file_lines:
            continue

        # 若启用单文件上限，截取前 N 行（用于后 30 页保证覆盖面）
        if max_lines_per_file and len(file_lines) > max_lines_per_file + 1:  # +1 保留分隔行
            file_lines = file_lines[:max_lines_per_file + 1]
            file_lines.append(f"/* ... {rel_path} 后续内容截断 ... */")

        # 预估加入后是否超限
        remaining = MAX_LINES - len(all_lines)
        if len(file_lines) > remaining:
            file_lines = file_lines[:remaining]

        all_lines.extend(file_lines)
        included_files.append(rel_path)

    return all_lines, included_files


# ============================================================
# Word 生成
# ============================================================
def set_cell_font(run, font_name: str = "Consolas", size_pt: int = 8):
    """设置代码字体。"""
    run.font.name = font_name
    run.font.size = Pt(size_pt)
    run.font.color.rgb = RGBColor(0, 0, 0)
    # 对中文字体做兼容性设置
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")


def add_field(run, field_code: str):
    """在 run 中插入 Word 域代码（如 PAGE / NUMPAGES）。"""
    fld_char_begin = OxmlElement("w:fldChar")
    fld_char_begin.set(qn("w:fldCharType"), "begin")

    instr_text = OxmlElement("w:instrText")
    instr_text.set(qn("xml:space"), "preserve")
    instr_text.text = field_code

    fld_char_end = OxmlElement("w:fldChar")
    fld_char_end.set(qn("w:fldCharType"), "end")

    run._r.append(fld_char_begin)
    run._r.append(instr_text)
    run._r.append(fld_char_end)


def set_section_start_page(section, start_num: int):
    """设置该节的起始页码（用于后 30 页从 31 开始）。"""
    sectPr = section._sectPr
    pgNumType = OxmlElement("w:pgNumType")
    pgNumType.set(qn("w:start"), str(start_num))
    sectPr.append(pgNumType)


def create_docx(
    lines: list[str],
    output_path: Path,
    included_files: list[str],
    *,
    start_page: int = 1,
    total_pages_label: str | None = None,
):
    """生成标准 A4 50 行/页的 docx。"""
    doc = Document()

    # 页面设置为 A4
    section = doc.sections[0]
    section.page_height = Inches(11.69)
    section.page_width = Inches(8.27)
    section.top_margin = Inches(0.6)
    section.bottom_margin = Inches(0.6)
    section.left_margin = Inches(0.8)
    section.right_margin = Inches(0.8)

    if start_page != 1:
        set_section_start_page(section, start_page)

    # 页眉
    header = section.header
    header_para = header.paragraphs[0]
    header_para.text = f"{SOFTWARE_NAME} {SOFTWARE_VERSION}"
    header_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in header_para.runs:
        run.font.name = "宋体"
        run.font.size = Pt(9)
        run._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")

    # 页脚：第 X 页 共 Y 页
    footer = section.footer
    footer_para = footer.paragraphs[0]
    footer_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_field(footer_para.add_run("第 "), "PAGE")
    footer_para.add_run(" 页 共 ")
    if total_pages_label:
        footer_para.add_run(total_pages_label)
    else:
        add_field(footer_para.add_run(""), "NUMPAGES")
    footer_para.add_run(" 页")
    for run in footer_para.runs:
        run.font.name = "宋体"
        run.font.size = Pt(9)
        run._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")

    total_pages = (len(lines) + LINES_PER_PAGE - 1) // LINES_PER_PAGE

    # 跟踪当前文件名，用于续接标记
    current_file: str | None = None

    for page_idx in range(total_pages):
        start = page_idx * LINES_PER_PAGE
        end = min(start + LINES_PER_PAGE, len(lines))
        page_lines = lines[start:end]

        # 判断下一页是否继续同一文件
        next_file = None
        if end < len(lines):
            next_line = lines[end]
            if next_line.startswith("/* ===== ") and next_line.endswith(" ===== */"):
                next_file = next_line[8:-8]
            else:
                next_file = current_file

        # 页面首行：若当前页第一行不是文件分隔行，且 current_file 已设置，则加续接标记
        first_is_separator = page_lines and page_lines[0].startswith("/* ===== ")
        if not first_is_separator and current_file:
            page_lines.insert(0, f"/* 续：{current_file} */")
            page_lines = page_lines[:LINES_PER_PAGE]  # 保持每页 50 行

        # 页面末行：若本页末尾未读完当前文件，且下一页继续同一文件，则加续标记
        last_is_separator = page_lines and page_lines[-1].startswith("/* ===== ")
        if not last_is_separator and next_file and next_file == current_file:
            if len(page_lines) == LINES_PER_PAGE:
                page_lines[-1] = "/* 续 */"
            else:
                page_lines.append("/* 续 */")

        # 写入页面内容
        for line in page_lines:
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(0)
            p.paragraph_format.space_after = Pt(0)
            p.paragraph_format.line_spacing = Pt(13.5)  # 固定行距，约 50 行/页
            p.paragraph_format.line_spacing_rule = None  # 使用绝对行距
            run = p.add_run(line)
            set_cell_font(run, "Consolas", 8)

            # 更新 current_file：遇到文件分隔行时更新
            if line.startswith("/* ===== ") and line.endswith(" ===== */"):
                current_file = line[8:-8]
            elif line.startswith("/* 续：") and line.endswith(" */"):
                current_file = line[5:-3]

        # 添加分页符（最后一页除外）
        if page_idx < total_pages - 1:
            doc.add_page_break()

    # 保存
    output_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(output_path)
    print(f"[OK] 已生成: {output_path}")
    print(f"     纳入文件数: {len(included_files)}")
    print(f"     总行数: {len(lines)}")
    print(f"     总页数: {total_pages}")


def generate_part(
    file_order: list[str],
    output_name: str,
    start_page: int,
    total_pages_label: str | None,
    max_lines_per_file: int | None = None,
):
    print(f"\n[INFO] 开始生成 {output_name} ...")
    lines, included_files = collect_content(file_order, max_lines_per_file)
    print(f"[INFO] 已收集 {len(lines)} 行，来自 {len(included_files)} 个文件")

    output_path = OUTPUT_DIR / output_name
    create_docx(lines, output_path, included_files, start_page=start_page, total_pages_label=total_pages_label)

    print("\n纳入文件清单:")
    for f in included_files:
        print(f"  - {f}")
    return included_files


def main():
    # 前置校验：确认 package.json 名称/版本
    package_json = PROJECT_ROOT / "package.json"
    import json
    with open(package_json, "r", encoding="utf-8") as f:
        pkg = json.load(f)
    if pkg.get("name") != "ai-film-flow" or pkg.get("version") != "1.0.0":
        print(
            f"[ERROR] package.json 校验失败: name={pkg.get('name')}, version={pkg.get('version')}，"
            "要求 name='ai-film-flow' 且 version='1.0.0'，已停止生成。"
        )
        sys.exit(1)
    print("[OK] package.json 校验通过: ai-film-flow v1.0.0")

    generate_part(
        FILE_ORDER_FRONT,
        "源代码前30页.docx",
        start_page=1,
        total_pages_label=None,  # 使用 NUMPAGES 域
        max_lines_per_file=None,
    )
    generate_part(
        FILE_ORDER_BACK,
        "源代码后30页.docx",
        start_page=31,
        total_pages_label="60",  # 合并后总页数固定 60
        max_lines_per_file=None,
    )


if __name__ == "__main__":
    main()
